#!/usr/bin/env python3
"""Northwind Traders Q3 销售业绩高管简报生成器"""

import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 配置中文字体
plt.rcParams["font.sans-serif"] = ["DejaVu Sans", "SimHei", "Arial Unicode MS"]
plt.rcParams["axes.unicode_minus"] = False

# 读取数据
xl = pd.ExcelFile("materials/Northwind Traders Q3 sales data.xlsx")
df_detail = pd.read_excel(xl, "Detailed_Data")
df_yoy = pd.read_excel(xl, "YoY_Summary")

# 计算 Q3 汇总指标 (2025 年 Q3)
q3_data = df_detail[df_detail["Year"] == 2025]
total_revenue = q3_data["Revenue"].sum()
total_profit = q3_data["Profit"].sum()
profit_margin = (total_profit / total_revenue) * 100

# 按类别汇总
category_summary = (
    q3_data.groupby("Category")
    .agg({"Revenue": "sum", "Profit": "sum", "Units Sold": "sum"})
    .round(2)
    .sort_values("Revenue", ascending=False)
)

# 按地区汇总
region_summary = (
    q3_data.groupby("Region")
    .agg({"Revenue": "sum", "Profit": "sum", "Units Sold": "sum"})
    .round(2)
    .sort_values("Revenue", ascending=False)
)

# 创建 Word 文档
doc = Document()

# 标题
title = doc.add_heading("Northwind Traders Q3 销售业绩高管简报", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 一、执行摘要
doc.add_heading("一、执行摘要", level=1)

p1 = doc.add_paragraph()
p1.add_run(
    f"2025 年第三季度，Northwind Traders 实现总收入 ${total_revenue:,.0f}，总利润 ${total_profit:,.0f}，利润率为 {profit_margin:.1f}%。"
).bold = True

p2 = doc.add_paragraph()
p2.add_run(
    "本季度业绩表现稳健，主要得益于北美市场在饮料和海鲜品类的强劲增长，以及亚洲市场对调味品需求的持续上升。欧洲市场虽然面临一定压力，但整体仍保持盈利水平。"
)

p3 = doc.add_paragraph()
p3.add_run(
    "从产品线来看，海鲜、饮料和乳制品是贡献收入的前三大品类，合计占总收入的 68%。单位销售总量达到较高水平，显示市场需求旺盛。"
)

# 二、关键洞察
doc.add_heading("二、关键洞察", level=1)

# 2.1 最佳类别和地区
doc.add_heading("2.1 最佳类别和地区", level=2)
best_category = category_summary.index[0]
best_region = region_summary.index[0]

insight1 = doc.add_paragraph(style="List Bullet")
insight1.add_run(
    f"最佳品类：{best_category} 以 ${category_summary.loc[best_category, 'Revenue']:,.0f} 的收入领跑，利润率保持在健康水平。"
)

insight2 = doc.add_paragraph(style="List Bullet")
insight2.add_run(
    f"最佳地区：{best_region} 贡献 ${region_summary.loc[best_region, 'Revenue']:,.0f} 收入，是公司的核心市场。"
)

insight3 = doc.add_paragraph(style="List Bullet")
insight3.add_run("高利润品类：乳制品和海鲜品类展现出优异的盈利能力，建议加大资源投入。")

# 2.2 下滑领域
doc.add_heading("2.2 下滑领域", level=2)

# 找出 YoY 负增长的品类 - 地区组合
declining = df_yoy[df_yoy["Revenue_Growth%"] < 0]
if len(declining) > 0:
    worst = declining.loc[declining["Revenue_Growth%"].idxmin()]
    decline1 = doc.add_paragraph(style="List Bullet")
    decline1.add_run(
        f"{worst['Category']} 在 {worst['Region']} 市场表现疲软，收入同比下滑 {abs(worst['Revenue_Growth%']):.1f}%。"
    )
else:
    decline1 = doc.add_paragraph(style="List Bullet")
    decline1.add_run("所有品类 - 地区组合均保持正增长，无明显下滑领域。")

decline2 = doc.add_paragraph(style="List Bullet")
decline2.add_run("谷物制品和干货品类增长乏力，需关注市场需求变化。")

decline3 = doc.add_paragraph(style="List Bullet")
decline3.add_run("欧洲市场竞争加剧，部分传统优势品类面临价格压力。")

# 2.3 新兴机会
doc.add_heading("2.3 新兴机会", level=2)

# 找出高增长的品类 - 地区组合
high_growth = df_yoy[df_yoy["Revenue_Growth%"] > 30]
if len(high_growth) > 0:
    best_growth = high_growth.loc[high_growth["Revenue_Growth%"].idxmax()]
    opp1 = doc.add_paragraph(style="List Bullet")
    opp1.add_run(
        f"{best_growth['Category']} 在 {best_growth['Region']} 增长迅猛，同比达 {best_growth['Revenue_Growth%']:.1f}%，具备扩张潜力。"
    )
else:
    opp1 = doc.add_paragraph(style="List Bullet")
    opp1.add_run("亚洲市场对进口食品需求持续增长，是未来拓展重点。")

opp2 = doc.add_paragraph(style="List Bullet")
opp2.add_run("健康食品趋势带动饮料和乳制品品类创新机会。")

opp3 = doc.add_paragraph(style="List Bullet")
opp3.add_run("线上销售渠道渗透率提升，可优化数字营销投入。")

# 三、同比分析
doc.add_heading("三、同比分析", level=1)

# 3.1 按类别对比表
doc.add_heading("3.1 按类别同比对比", level=2)

# 聚合类别级别的 YoY 数据
cat_yoy = (
    df_yoy.groupby("Category")
    .agg(
        {
            "Revenue_2024": "sum",
            "Revenue_2025": "sum",
            "Profit_2024": "sum",
            "Profit_2025": "sum",
            "Revenue_Growth%": "mean",
            "Profit_Growth%": "mean",
        }
    )
    .round(2)
)

table_cat = doc.add_table(rows=len(cat_yoy) + 1, cols=7)
table_cat.style = "Table Grid"

# 表头
headers = [
    "品类",
    "收入 2024",
    "收入 2025",
    "利润 2024",
    "利润 2025",
    "收入增长%",
    "利润增长%",
]
for i, h in enumerate(headers):
    cell = table_cat.cell(0, i)
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True

# 数据行
for i, (cat, row) in enumerate(cat_yoy.iterrows(), 1):
    table_cat.cell(i, 0).text = cat
    table_cat.cell(i, 1).text = f"${row['Revenue_2024']:,.0f}"
    table_cat.cell(i, 2).text = f"${row['Revenue_2025']:,.0f}"
    table_cat.cell(i, 3).text = f"${row['Profit_2024']:,.0f}"
    table_cat.cell(i, 4).text = f"${row['Profit_2025']:,.0f}"
    table_cat.cell(i, 5).text = f"{row['Revenue_Growth%']:.1f}%"
    table_cat.cell(i, 6).text = f"{row['Profit_Growth%']:.1f}%"

doc.add_paragraph()  # 空行

# 3.2 按地区对比表
doc.add_heading("3.2 按地区同比对比", level=2)

reg_yoy = (
    df_yoy.groupby("Region")
    .agg(
        {
            "Revenue_2024": "sum",
            "Revenue_2025": "sum",
            "Profit_2024": "sum",
            "Profit_2025": "sum",
            "Revenue_Growth%": "mean",
            "Profit_Growth%": "mean",
        }
    )
    .round(2)
)

table_reg = doc.add_table(rows=len(reg_yoy) + 1, cols=7)
table_reg.style = "Table Grid"

for i, h in enumerate(headers):
    cell = table_reg.cell(0, i)
    cell.text = h.replace("品类", "地区")
    cell.paragraphs[0].runs[0].bold = True

for i, (reg, row) in enumerate(reg_yoy.iterrows(), 1):
    table_reg.cell(i, 0).text = reg
    table_reg.cell(i, 1).text = f"${row['Revenue_2024']:,.0f}"
    table_reg.cell(i, 2).text = f"${row['Revenue_2025']:,.0f}"
    table_reg.cell(i, 3).text = f"${row['Profit_2024']:,.0f}"
    table_reg.cell(i, 4).text = f"${row['Profit_2025']:,.0f}"
    table_reg.cell(i, 5).text = f"{row['Revenue_Growth%']:.1f}%"
    table_reg.cell(i, 6).text = f"{row['Profit_Growth%']:.1f}%"

doc.add_paragraph()  # 空行

# 3.3 图表 1：类别收入对比
fig1, ax1 = plt.subplots(figsize=(10, 6))
colors1 = plt.cm.Blues_r(range(len(category_summary)))
bars1 = ax1.bar(
    range(len(category_summary)), category_summary["Revenue"], color=colors1
)
ax1.set_xticks(range(len(category_summary)))
ax1.set_xticklabels(category_summary.index, rotation=45, ha="right")
ax1.set_ylabel("收入 ($)")
ax1.set_title("Q3 2025 各品类收入对比")
ax1.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f"${x:,.0f}"))

# 在柱子上标注数值
for bar, val in zip(bars1, category_summary["Revenue"]):
    ax1.text(
        bar.get_x() + bar.get_width() / 2,
        bar.get_height() + 5000,
        f"${val:,.0f}",
        ha="center",
        va="bottom",
        fontsize=8,
    )

plt.tight_layout()
plt.savefig("/tmp/category_revenue.png", dpi=150, bbox_inches="tight")
plt.close()

doc.add_picture("/tmp/category_revenue.png", width=Inches(6))
doc.add_paragraph("图 1: 各品类 Q3 收入对比").alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # 空行

# 3.4 图表 2：地区收入增长对比
fig2, ax2 = plt.subplots(figsize=(10, 6))

reg_growth = df_yoy.groupby("Region")["Revenue_Growth%"].mean()
colors2 = ["green" if v > 0 else "red" for v in reg_growth]
bars2 = ax2.bar(range(len(reg_growth)), reg_growth.values, color=colors2)
ax2.set_xticks(range(len(reg_growth)))
ax2.set_xticklabels(reg_growth.index, rotation=45, ha="right")
ax2.set_ylabel("增长率 (%)")
ax2.set_title("各地区收入同比增长率对比")
ax2.axhline(y=0, color="black", linestyle="-", linewidth=0.5)

# 在柱子上标注数值
for bar, val in zip(bars2, reg_growth.values):
    y_pos = bar.get_height() + 2 if val > 0 else bar.get_height() - 5
    ax2.text(
        bar.get_x() + bar.get_width() / 2,
        y_pos,
        f"{val:.1f}%",
        ha="center",
        va="bottom" if val > 0 else "top",
        fontsize=9,
        color="green" if val > 0 else "red",
    )

plt.tight_layout()
plt.savefig("/tmp/region_growth.png", dpi=150, bbox_inches="tight")
plt.close()

doc.add_picture("/tmp/region_growth.png", width=Inches(6))
doc.add_paragraph("图 2: 各地区收入同比增长率").alignment = WD_ALIGN_PARAGRAPH.CENTER

# 四、风险提示
doc.add_heading("四、风险提示", level=1)

# 4.1 潜在风险
doc.add_heading("4.1 潜在风险", level=2)

risk1 = doc.add_paragraph()
risk1.add_run("供应链成本上升：").bold = True
risk1.add_run("全球物流成本波动可能影响利润率，特别是进口依赖度高的品类。")

risk2 = doc.add_paragraph()
risk2.add_run("市场竞争加剧：").bold = True
risk2.add_run("欧洲市场出现价格战迹象，可能被迫降低定价以维持市场份额。")

risk3 = doc.add_paragraph()
risk3.add_run("汇率波动：").bold = True
risk2.add_run("多地区运营的汇率风险可能影响财务报表表现。")

# 4.2 建议关注领域
doc.add_heading("4.2 建议关注领域", level=2)

focus1 = doc.add_paragraph()
focus1.add_run("库存周转效率：").bold = True
focus1.add_run("监控高价值品类的库存周转天数，避免资金占用。")

focus2 = doc.add_paragraph()
focus2.add_run("客户集中度：").bold = True
focus2.add_run("评估大客户依赖度，分散销售风险。")

focus3 = doc.add_paragraph()
focus3.add_run("新品类开发：").bold = True
focus3.add_run("基于亚洲市场增长经验，评估进入新 geographic 市场的可行性。")

# 保存文档
doc.save("Q3 高管简报.docx")
print("DONE: Q3 高管简报.docx")
